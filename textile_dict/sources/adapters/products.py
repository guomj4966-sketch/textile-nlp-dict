"""专精特新产品/企业案例适配器 — 从企业产品目录和案例文档提取术语。

数据源类型: product
数据源位置:
  - Excel: ../专精特新平台/新产品（文字）/2024-2026产品技术目录汇总.xlsx
  - Excel: ../专精特新平台/前五批纺织行业专精特新企业汇总（培育入库）.xlsx
  - docx:  ../专精特新平台/企业案例/
  - docx:  ../专精特新平台/新产品（文字）/
"""

from __future__ import annotations

import re
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Iterator, Optional

from textile_dict.sources.adapter import (
    DataSourceAdapter,
    ExtractionResult,
    SourceMetadata,
    TermCandidate,
)
from textile_dict.sources.shared import (
    classify_candidate,
    discover_terms_from_texts,
    is_textile_relevant,
)

# 数据源根目录
_ONEDRIVE = Path(__file__).parent.parent.parent.parent
PRODUCTS_BASE = _ONEDRIVE / "专精特新平台"

# 已知的 Excel 和目录
_PRODUCT_DIRS: list[tuple[Path, str]] = [
    (PRODUCTS_BASE / "新产品（文字）" / "2024-2026产品技术目录汇总.xlsx", "产品目录汇总"),
    (PRODUCTS_BASE / "前五批纺织行业专精特新企业汇总（培育入库）.xlsx", "前五批特色产品"),
]
_CASE_DIRS: list[Path] = [
    PRODUCTS_BASE / "企业案例" / "企业案例-2022发展报告",
    PRODUCTS_BASE / "企业案例" / "企业案例-2026绿色发展",
    PRODUCTS_BASE / "新产品（文字）" / "2024汇编材料文字版",
    PRODUCTS_BASE / "新产品（文字）" / "2025汇编材料文字版",
    PRODUCTS_BASE / "新产品（文字）" / "2026绿色产品技术",
]


class ProductSource(DataSourceAdapter):
    """专精特新企业案例和产品资料适配器。

    从 Excel 产品名 + docx 企业案例 + PDF 研究报告中提取纺织术语。
    """

    name = "products"

    def __init__(
        self,
        base_dir: Optional[str | Path] = None,
    ):
        self._base = Path(base_dir) if base_dir else PRODUCTS_BASE
        self._meta: Optional[SourceMetadata] = None

    # ── 元信息 ──

    @property
    def metadata(self) -> SourceMetadata:
        if self._meta is not None:
            return self._meta

        exists = self._base.exists()
        item_count = 0
        if exists:
            # 粗略统计
            item_count = len(list(self._base.rglob("*.docx")))
            item_count += len(list(self._base.rglob("*.xlsx")))
            item_count += len(list(self._base.rglob("*.pdf")))

        self._meta = SourceMetadata(
            name="products",
            display_name="专精特新企业案例与产品库",
            source_type="product",
            version="1.0",
            description="纺织行业专精特新企业案例、产品技术目录、特色产品表、研究报告",
            location=str(self._base),
            item_count=item_count,
            tags=["纺织", "专精特新", "企业案例", "新产品", "产品目录"],
        )
        return self._meta

    # ── 连接 ──

    def validate_connection(self) -> bool:
        return self._base.exists()

    # ── 文本迭代 ──

    def iter_texts(self) -> Iterator[str]:
        """从 Excel + docx + PDF 中逐条产出全文。

        每种格式一个 yield。Excel 产出产品名拼接文，
        docx/PDF 产出完整段落。
        """
        yield from self._iter_excel_product_names()
        yield from self._iter_excel_characteristic_products()
        yield from self._iter_docx_files()
        yield from self._iter_pdf_files()

    # ── Excel 产品名 ──

    def _iter_excel_product_names(self) -> Iterator[str]:
        path = self._base / "新产品（文字）" / "2024-2026产品技术目录汇总.xlsx"
        if not path.exists():
            return
        try:
            from openpyxl import load_workbook
        except ImportError:
            return

        wb = load_workbook(path, data_only=True)
        names: list[str] = []

        if "汇总" in wb.sheetnames:
            ws = wb["汇总"]
            for row in ws.iter_rows(min_row=2, values_only=True):
                if row and len(row) >= 2 and row[1]:
                    names.append(str(row[1]).strip())

        for sn in ["2024", "2025", "2026"]:
            if sn in wb.sheetnames:
                ws = wb[sn]
                for row in ws.iter_rows(min_row=1, max_row=10, values_only=True):
                    if row and any("编号" in str(c) for c in row if c):
                        break
                for row in ws.iter_rows(min_row=2, values_only=True):
                    if row and len(row) >= 3:
                        name = str(row[2] if len(row) > 2 else row[1]).strip()
                        if name and len(name) > 3:
                            names.append(name)
        wb.close()

        if names:
            # 分成多批 yield（避免单条文本太大）
            for i in range(0, len(names), 50):
                batch = names[i : i + 50]
                yield "产品名称：" + "；".join(batch)

    def _iter_excel_characteristic_products(self) -> Iterator[str]:
        path = self._base / "前五批纺织行业专精特新企业汇总（培育入库）.xlsx"
        if not path.exists():
            return
        try:
            import pandas as pd
        except ImportError:
            return

        xl = pd.ExcelFile(path)
        if "前5批特色产品" not in xl.sheet_names:
            xl.close()
            return

        df = pd.read_excel(path, sheet_name="前5批特色产品")
        for col in ["特色产品名称", "特色产品介绍"]:
            if col in df.columns:
                items = df[col].dropna().tolist()
                for item in items:
                    s = str(item).replace("_x000D_", "").strip()
                    if len(s) > 3:
                        yield s

    # ── docx ──

    def _iter_docx_files(self) -> Iterator[str]:
        try:
            from docx import Document
        except ImportError:
            return

        for case_dir in _CASE_DIRS:
            if not case_dir.exists():
                continue
            for fp in case_dir.glob("*.docx"):
                try:
                    doc = Document(str(fp))
                    paragraphs = [
                        p.text.strip() for p in doc.paragraphs if p.text.strip()
                    ]
                    if paragraphs:
                        yield fp.stem + "。" + "\n".join(paragraphs)
                except Exception:
                    pass
            # Also .doc
            for fp in case_dir.glob("*.doc"):
                try:
                    doc = Document(str(fp))
                    paragraphs = [
                        p.text.strip() for p in doc.paragraphs if p.text.strip()
                    ]
                    if paragraphs:
                        yield fp.stem + "。" + "\n".join(paragraphs)
                except Exception:
                    pass

    # ── PDF ──

    def _iter_pdf_files(self) -> Iterator[str]:
        try:
            import pdfplumber
        except ImportError:
            return

        report_dir = self._base / "研究报告"
        if not report_dir.exists():
            return

        for fp in report_dir.glob("*.pdf"):
            total_chars = 0
            try:
                with pdfplumber.open(str(fp)) as pdf:
                    for page in pdf.pages:
                        txt = page.extract_text() or ""
                        total_chars += len(txt)
                    if total_chars > 500:
                        full_texts: list[str] = []
                        for page in pdf.pages:
                            t = page.extract_text() or ""
                            if t.strip():
                                full_texts.append(t)
                        yield fp.stem + "。" + "\n".join(full_texts)
            except Exception:
                pass

    # ── 术语提取 ──

    def extract_terms(
        self,
        texts: Optional[list[str]] = None,
        min_len: int = 2,
        min_freq: int = 2,
        max_terms: int = 500,
    ) -> ExtractionResult:
        started_at = datetime.now().isoformat()

        if texts is None:
            texts = list(self.iter_texts())

        total_chars = sum(len(t) for t in texts)

        # 使用共享提取引擎
        discovered = discover_terms_from_texts(
            texts,
            min_len=min_len,
            min_freq=min_freq,
            max_terms=max_terms * 2,  # 先多取，后续按纺织相关性过滤
        )

        candidates: list[TermCandidate] = []
        for term, freq in discovered.most_common():
            if len(candidates) >= max_terms:
                break

            # 产品源的术语都应与纺织相关（因为专精特新筛选后的企业都是纺织相关）
            # 但再做一遍织物相关性检查
            if not is_textile_relevant(term):
                continue

            layer, category = classify_candidate(term)
            confidence = min(1.0, 0.35 + 0.05 * freq)

            candidates.append(
                TermCandidate(
                    term=term,
                    source_name=self.name,
                    layer=layer,
                    category=category,
                    frequency=freq,
                    confidence_score=round(confidence, 2),
                )
            )

        return ExtractionResult(
            adapter_name=self.name,
            extracted_at=started_at,
            total_texts_processed=len(texts),
            total_chars_processed=total_chars,
            candidates=candidates,
            stats={
                "raw_discovered": len(discovered),
                "textile_relevant": len(candidates),
            },
        )
